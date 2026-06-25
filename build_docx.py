#!/usr/bin/env python3
"""
Build AER-MQoS.docx from LaTeX sources + figures, matching IEEE Word template style.
Usage: python3 build_docx.py
Output: ../AER-MQoS.docx
"""
import re, sys
from pathlib import Path
from copy import deepcopy

import docx
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

SRC = Path(__file__).resolve().parent / "Section-tex"
FIGS = SRC / "Figures"
OUT = Path(__file__).resolve().parent / "AER-MQoS.docx"

SECTIONS = [
    "sections/abstract",
    "sections/introduction",
    "sections/related_work",
    "sections/architecture",
    "sections/security",
    "sections/energy",
    "sections/qlearning",
    "sections/evaluation",
    "sections/conclusion",
]

# Additional sections after main text (Data Availability, Author contributions)
BONUS_SECTIONS = [
    "sections/appendix",
]

FIGURE_MAP = {
    "Fig_1_Architecture_AER_MQoS": "Figure 1",
    "Fig_2_DODAG_Path_Levels": "Figure 2",
    "Fig_3_Context_Weights_Alpha_Beta_Gamma": "Figure 3",
    "Fig_4_PDR_Comparison_4_Protocols": "Figure 4",
    "Fig_5_Latency_Comparison_4_Protocols": "Figure 5",
    "Fig_6_Interarrival_Gap_By_Traffic_Class": "Figure 6",
    "Fig_7_Energy_Proxy_Comparison": "Figure 7",
    "Fig_8_Control_Overhead_RPL_Messages": "Figure 8",
    "Fig_9_Convergence_Or_Stability_Time": "Figure 9",
    "Fig_10_Temporal_PDR_Stratification": "Figure 10",
    "Fig_11_Learning_Or_Load_Sensitivity": "Figure 11",
}


def _fmt(run, text, bold=False, italic=False, mono=False, size=None, color=None):
    run.text = text
    run.bold = bold
    run.italic = italic
    if mono:
        run.font.name = "Courier New"
        run.font.size = Pt(size or 9)
    else:
        run.font.name = "Times New Roman"
        if size:
            run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text, style="Normal", bold=False, italic=False, space_after=Pt(6), alignment=None):
    p = doc.add_paragraph(style=style)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p


def add_rich_para(doc, parts, style="Normal", space_after=Pt(6), alignment=None):
    """parts = list of (text, bold, italic, mono, size) tuples"""
    p = doc.add_paragraph(style=style)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    for text, bold, italic, mono, size in parts:
        run = p.add_run(text)
        run.font.name = "Courier New" if mono else "Times New Roman"
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
    return p


def strip_latex(text):
    """Remove LaTeX commands, keep readable text."""
    # Remove comments
    text = re.sub(r'(?<!\\)%.*', '', text)
    # Remove \label{...}
    text = re.sub(r'\\label\{[^}]*\}', '', text)
    # Remove \ref{...} and \cite{...} - keep placeholder
    text = re.sub(r'\\(ref|cite)\{[^}]*\}', '[REF]', text)
    # Remove \textsf, \texttt, \textbf, \textit wrappers - keep content
    text = re.sub(r'\\(textsf|texttt|textbf|textit|emph)\{([^}]*)\}', r'\2', text)
    # Remove \AERMQoS macro
    text = re.sub(r'\\AERMQoS', 'AER-MQoS', text)
    # Remove \PaperAuthorNameShort{...}
    text = re.sub(r'\\PaperAuthorNameShort\{[^}]*\}', 'Madani Belacel', text)
    # Remove \PaperTitle etc.
    text = re.sub(r'\\(PaperTitle|PaperAuthorName|PaperEmail|PaperCorrespondingEmail|AuthorOrcid)[^}]*\}', '', text)
    # Handle \texorpdfstring{tex}{pdf} → keep pdf (plain-text) arg
    text = re.sub(r'\\texorpdfstring\{([^}]*)\}\{([^}]*)\}', r'\2', text)
    # Greek letters → Unicode
    GREEK = {
        'alpha': 'α', 'beta': 'β', 'gamma': 'γ', 'delta': 'δ',
        'theta': 'θ', 'lambda': 'λ', 'mu': 'μ', 'sigma': 'σ',
        'omega': 'ω', 'phi': 'φ', 'pi': 'π', 'tau': 'τ',
        'varepsilon': 'ε', 'rho': 'ρ', 'eta': 'η',
    }
    for cmd, char in GREEK.items():
        text = re.sub(rf'\\{cmd}(?![a-zA-Z])', char, text)
    # Remove other macros
    text = re.sub(r'\\[a-zA-Z]+(\{[^}]*\})?', '', text)
    # Clean up $...$ math
    text = re.sub(r'\$(.*?)\$', r'\1', text)
    # Clean up ~ and --
    text = text.replace('~', ' ')
    text = text.replace('--', '–')
    text = text.replace('---', '—')
    # Collapse whitespace
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n\s*\n', '\n\n', text)
    return text.strip()


def read_section(filepath):
    """Read a .tex section file and return raw text."""
    # Try with .tex extension
    path = SRC / f"{filepath}.tex"
    if not path.exists():
        path = SRC / filepath
    if not path.exists():
        print(f"WARN: {path} not found", file=sys.stderr)
        return ""
    return path.read_text(encoding="utf-8")


def extract_figure_path(fig_name):
    """Return the path for a figure (PNG preferred for Word), or None."""
    png = FIGS / f"{fig_name}.png"
    pdf = FIGS / f"{fig_name}.pdf"
    if png.exists():
        return png
    if pdf.exists():
        return pdf
    return None


# Pre-defined figure caption text (from CAPTIONS_EN.tex)
FIGURE_CAPTIONS = {
    "Fig_1_Architecture_AER_MQoS":
        "Conceptual schematic—not simulation output. Reference architecture of AER-MQoS: "
        "QoS plane, trust/energy context plane, MCS, and experimental OCP 8 objective function.",
    "Fig_2_DODAG_Path_Levels":
        "Conceptual schematic—not simulation output. Illustrative RPL DODAG "
        "(not the 25-node Cooja layout): DIO (down), DAO (up), preferred parent path.",
    "Fig_3_Context_Weights_Alpha_Beta_Gamma":
        "Context coefficient γ (α=γ, β=100-γ) per traffic class C0–C3 from context.csv.",
    "Fig_4_PDR_Comparison_4_Protocols":
        "Overall and C3 PDR (%) on the lossy UDGM channel: per-seed boxplots and "
        "descriptive means from pdr.csv. Campaign descriptive means (Table VII); "
        "per-seed variance unavailable. RPL_MQoS and RPL_AER: 1 valid seed; "
        "AER-MQoS: seeds 2–4.",
    "Fig_5_Latency_Comparison_4_Protocols":
        "Measured application-layer end-to-end latency (ms) on the lossy channel. "
        "Campaign descriptive means (Table VII); per-seed variance unavailable.",
    "Fig_6_Interarrival_Gap_By_Traffic_Class":
        "Mean inter-arrival gap (ms) per class C0–C3 from jitter.csv "
        "(not RFC 3393 IPDV; C3 ring-highlighted). Campaign descriptive means (Table VII).",
    "Fig_7_Energy_Proxy_Comparison":
        "Internal NRE (instrumentation) per firmware tag from energy.csv "
        "(nre_proxy_pct); not Energest duty-cycle or joules.",
    "Fig_8_Control_Overhead_RPL_Messages":
        "Control-plane export proxy: ctrl_export_rate from ctrl.csv. "
        "Time axis covers the initial bootstrap window (≈0–1 min of 30 min simulation).",
    "Fig_9_Convergence_Or_Stability_Time":
        "Route-adaptation proxy: cumulative ctx_update_cumulative from stab.csv. "
        "Time axis covers the initial bootstrap window (≈0–1 min of 30 min simulation).",
    "Fig_10_Temporal_PDR_Stratification":
        "Temporal PDR stratification over epochs from sec.csv (four firmware tags).",
    "Fig_11_Learning_Or_Load_Sensitivity":
        "Logging sanity check: PDR vs. offered-load quartiles for learning_on strata "
        "in learn_or_load.csv (same AER-MQoS binary; not a learning ablation). "
        "Campaign descriptive means (Table VII); per-seed variance unavailable.",
}


def extract_caption(env_text, fig_label):
    """Extract caption from a figure environment, fall back to pre-defined text."""
    cap_match = re.search(r'\\caption\s*\{(.*?)\}', env_text, re.DOTALL)
    if cap_match:
        cap_raw = cap_match.group(1)
        cap = strip_latex(cap_raw).strip()
        if cap:
            return cap
    # Fall back to pre-defined caption keyed by fig_name (last part of label)
    for fig_name, cap in FIGURE_CAPTIONS.items():
        if FIGURE_MAP.get(fig_name) == fig_label:
            return cap
    return fig_label


def is_header_line(line):
    """Check if line starts a section or subsection."""
    return bool(re.match(r'\\(section|subsection|subsubsection)\*?\{', line.strip()))


def parse_content(text):
    """Very simple LaTeX parser: yield (type, data) tokens."""
    lines = text.split('\n')
    in_env = None
    env_lines = []
    for line in lines:
        stripped = line.strip()

        if in_env:
            env_lines.append(line)
            if stripped.startswith(f"\\end{{{in_env}}}"):
                yield (f"env:{in_env}", env_lines)
                env_lines = []
                in_env = None
            continue

        if stripped.startswith("\\begin{"):
            m = re.match(r'\\begin\{(\w+)\}', stripped)
            if m:
                in_env = m.group(1)
                env_lines = [line]
                continue

        if is_header_line(line):
            yield ("header", line)
        elif stripped.startswith("\\includegraphics"):
            yield ("graphics", line)
        elif stripped.startswith("\\caption"):
            yield ("caption", line)
        elif stripped:
            yield ("text", line)


def extract_balanced_braces(text, start):
    """Return (content, end_pos) of balanced braces starting at text[start]=='{'."""
    if start >= len(text) or text[start] != '{':
        return None, start
    depth = 1
    i = start + 1
    while i < len(text) and depth > 0:
        if text[i] == '{':
            depth += 1
        elif text[i] == '}':
            depth -= 1
        i += 1
    return text[start + 1 : i - 1], i


def parse_heading(line):
    """Return (level, title) from a section or subsection line."""
    m = re.match(r'\\(section|subsection|subsubsection)\*?\{', line)
    if not m:
        return None
    cmd = m.group(1)
    level = {'section': 1, 'subsection': 2, 'subsubsection': 3}[cmd]
    content, _ = extract_balanced_braces(line, m.end() - 1)
    return (level, strip_latex(content))


def parse_includegraphics(line):
    m = re.search(r'\{([^}]+)\}', line)
    if m:
        name = m.group(1)
        # Remove extension if present
        name = name.replace('.pdf', '').replace('.png', '')
        return name
    return None


def parse_table(env_lines):
    """Parse tabular/tabularx environment, return (headers: list, rows: list[list])."""
    text = '\n'.join(env_lines)
    # Remove \toprule, \midrule, \bottomrule
    text = re.sub(r'\\(toprule|midrule|bottomrule)', '', text)
    # Remove \hline
    text = re.sub(r'\\hline', '', text)
    # Extract rows between \\ and \\ or end
    rows_raw = re.split(r'\\\\', text)
    headers = []
    rows = []
    for row_text in rows_raw:
        if not row_text.strip():
            continue
        # Extract cells separated by &
        cells = re.findall(r'\{(.*?)\}', row_text)
        if not cells:
            cells = [c.strip() for c in row_text.replace('\n', ' ').split('&') if c.strip()]
        cells = [strip_latex(c) for c in cells]
        if not cells:
            continue
        if not headers:
            headers = cells
        else:
            rows.append(cells)
    return headers, rows


def build_document():
    doc = Document()

    # ── Page setup (letter, IEEE margins) ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.space_before = Pt(0)

    # ── Title ──
    add_para(doc, "AER-MQoS: A Context-Aware Multi-Objective RPL Extension for\n"
                  "QoS-Aware, Energy-Aware, and Link-Reliability-Informed Routing\n"
                  "in Low-Power IoT Networks",
             style="Normal", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

    # ── Authors ──
    add_para(doc, "Madani Belacel",
             style="Normal", italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    # ── Abstract ──
    abstract_text = read_section("sections/abstract")
    abstract_clean = strip_latex(abstract_text)
    # Remove the \section* if present
    abstract_clean = re.sub(r'\\begin\{abstract\}', '', abstract_clean)
    abstract_clean = re.sub(r'\\end\{abstract\}', '', abstract_clean)
    abstract_clean = abstract_clean.strip()

    # Split into abstract label + content
    add_para(doc, f"Abstract—{abstract_clean}", style="Normal", space_after=Pt(12))

    # ── Keywords ──
    add_para(doc, "Index Terms—RPL, IoT, low-power networks, QoS, energy-aware routing, "
                  "link-reliability routing, network simulation.",
             style="Normal", italic=True, space_after=Pt(12))

    # ── Process each section ──
    for sec_file in SECTIONS:
        if sec_file == "sections/abstract":
            continue  # already done
        sec_text = read_section(sec_file)
        if not sec_text:
            continue

        # Split into paragraphs and environments
        for token_type, data in parse_content(sec_text):
            if token_type == "header":
                level_title = parse_heading(data)
                if level_title:
                    level, title = level_title
                    if level == 1:
                        add_para(doc, title, style="Normal", bold=True, space_after=Pt(6))
                    elif level == 2:
                        add_para(doc, title, style="Normal", bold=True, italic=True, space_after=Pt(4))

            elif token_type == "text":
                clean = strip_latex(data)
                # Skip empty/header-only lines
                if clean and not is_header_line(data) and not clean.startswith('\\'):
                    # Process inline formatting
                    add_para(doc, clean)

            elif token_type == "graphics":
                fig_name = parse_includegraphics(data)
                if fig_name and fig_name in FIGURE_MAP:
                    fig_path = extract_figure_path(fig_name)
                    fig_label = FIGURE_MAP[fig_name]
                    if fig_path:
                        try:
                            doc.add_picture(str(fig_path), width=Inches(5.5))
                            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            # Use pre-defined caption text
                            cap_text = FIGURE_CAPTIONS.get(fig_name, fig_label)
                            add_para(doc, f"[{cap_text}]", style="Normal",
                                     italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                     space_after=Pt(10))
                        except Exception as e:
                            print(f"WARN: skipping {fig_path}: {e}", file=sys.stderr)
                else:
                    print(f"WARN: unknown figure '{fig_name}'", file=sys.stderr)

            elif token_type == "caption":
                pass  # handled within figure env

            elif token_type.startswith("env:"):
                env_name = token_type.split(":", 1)[1]
                env_text = '\n'.join(data)

                if env_name == "itemize":
                    # Extract items
                    items = re.findall(r'\\item\s+(.*?)(?=\\item|\Z)', env_text, re.DOTALL)
                    for item in items:
                        clean_item = strip_latex(item)
                        if clean_item:
                            add_para(doc, f"  \u2022  {clean_item}", space_after=Pt(2))

                elif env_name in ("table", "table*"):
                    # Find tabular inside
                    tab_match = re.search(r'\\begin\{tabular[x]*\}.*?\\end\{tabular[x]*\}',
                                          env_text, re.DOTALL)
                    if tab_match:
                        headers, rows = parse_table(tab_match.group())
                        # Create Word table
                        if rows:
                            table = doc.add_table(rows=len(rows) + 1, cols=max(len(headers), len(rows[0]) if rows else 1))
                            table.style = 'Table Grid'
                            # Headers
                            for j, h in enumerate(headers):
                                cell = table.rows[0].cells[j]
                                cell.text = h
                                for run in cell.paragraphs[0].runs:
                                    run.bold = True
                            # Data rows
                            for i, row_data in enumerate(rows):
                                for j, val in enumerate(row_data):
                                    if j < len(table.rows[i + 1].cells):
                                        table.rows[i + 1].cells[j].text = val

                elif env_name == "equation":
                    eq_text = strip_latex(env_text)
                    if eq_text:
                        add_para(doc, eq_text, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

                elif env_name in ("aertablecol", "aertablewide"):
                    # Extract tabular + caption
                    tab_match = re.search(r'\\begin\{tabular[x]*\}.*?\\end\{tabular[x]*\}',
                                          env_text, re.DOTALL)
                    if tab_match:
                        headers, rows = parse_table(tab_match.group())
                        if rows:
                            ncols = max(len(headers), len(rows[0])) if rows else 1
                            table = doc.add_table(rows=len(rows) + 1, cols=ncols)
                            table.style = 'Table Grid'
                            for j, h in enumerate(headers):
                                if j < len(table.rows[0].cells):
                                    table.rows[0].cells[j].text = h
                                    for run in table.rows[0].cells[j].paragraphs[0].runs:
                                        run.bold = True
                            for i, row_data in enumerate(rows):
                                for j, val in enumerate(row_data):
                                    if j < len(table.rows[i + 1].cells):
                                        table.rows[i + 1].cells[j].text = val
                    cap_match = re.search(r'\\caption\{([^}]*)\}', env_text)
                    if cap_match:
                        cap_text = strip_latex(cap_match.group(1))
                        add_para(doc, cap_text, style="Normal", italic=True, space_after=Pt(8))

                elif env_name in ("figure", "figure*"):
                    fig_name = None
                    fig_match = re.search(r'\\includegraphics\s*(?:\[[^\]]*\])?\s*\{([^}]+)\}', env_text)
                    if fig_match:
                        fig_name = fig_match.group(1).replace('.pdf', '').replace('.png', '')
                        fig_path = extract_figure_path(fig_name)
                        if fig_path:
                            try:
                                doc.add_picture(str(fig_path), width=Inches(5.5))
                                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except Exception as e:
                                print(f"WARN: skipping {fig_path}: {e}", file=sys.stderr)
                    fig_label = FIGURE_MAP.get(fig_name) if fig_name else None
                    cap_text = extract_caption(env_text, fig_label)
                    if cap_text:
                        add_para(doc, f"[{cap_text}]", style="Normal",
                                 italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                 space_after=Pt(10))

    # ── Data Availability ──
    add_para(doc, "Data Availability", bold=True, space_after=Pt(6))
    add_para(doc, "The source code, Cooja simulation scenarios, and measured datasets "
             "(N=4 independent seeds) used in this study are publicly available in the "
             "following GitHub repository:", space_after=Pt(4))
    add_para(doc, "Repository: https://github.com/madani-belacel/AER-MQoS", space_after=Pt(4))
    add_para(doc, "Contains: Contiki-NG firmware, simulation scripts, raw CSV data, "
             "and reproduction instructions.", space_after=Pt(8))

    # ── Author contributions ──
    add_para(doc, "Author contributions", bold=True, space_after=Pt(6))
    add_para(doc, "Madani Belacel conceived the integrated protocol design, implemented "
             "the Contiki-NG firmware and evaluation scripts, coordinated the Cooja campaigns, "
             "and wrote the manuscript.", space_after=Pt(8))

    # ── Appendix ──
    appendix_text = read_section("sections/appendix")
    if appendix_text:
        # Split into paragraphs and process
        for token_type, data in parse_content(appendix_text):
            if token_type == "header":
                level_title = parse_heading(data)
                if level_title:
                    level, title = level_title
                    add_para(doc, title, bold=True, space_after=Pt(6))
            elif token_type == "text":
                clean = strip_latex(data)
                if clean:
                    add_para(doc, clean)
            elif token_type.startswith("env:"):
                env_name = token_type.split(":", 1)[1]
                env_text = '\n'.join(data)
                if env_name in ("figure", "figure*"):
                    fig_match = re.search(r'\\includegraphics\s*(?:\[[^\]]*\])?\s*\{([^}]+)\}', env_text)
                    fig_label = None
                    if fig_match:
                        fig_name = fig_match.group(1).replace('.pdf', '').replace('.png', '')
                        fig_label = FIGURE_MAP.get(fig_name, fig_name)
                        fig_path = extract_figure_path(fig_name)
                        if fig_path:
                            try:
                                doc.add_picture(str(fig_path), width=Inches(5.5))
                                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except Exception as e:
                                print(f"WARN: skipping {fig_path}: {e}", file=sys.stderr)
                    cap_text = extract_caption(env_text, fig_label) if fig_label else None
                    if cap_text:
                        add_para(doc, f"[{cap_text}]", style="Normal",
                                 italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                 space_after=Pt(10))

    # ── Bibliography as plain text ──
    add_para(doc, "REFERENCES", style="Normal", bold=True, space_after=Pt(6))
    bib_path = SRC / "main-ieee.bbl"
    if bib_path.exists():
        bib_text = bib_path.read_text(encoding="utf-8")
        # Extract each \bibitem
        bib_entries = re.findall(r'\\bibitem\{[^}]*\}\s*(.*?)(?=\\bibitem|\Z)', bib_text, re.DOTALL)
        for i, entry in enumerate(bib_entries):
            clean = strip_latex(entry)
            if clean:
                add_para(doc, f"[{i + 1}] {clean}", style="Normal", space_after=Pt(4))
    else:
        add_para(doc, "[Bibliography not available; run BibTeX first]", italic=True)

    # ── Save ──
    doc.save(str(OUT))
    print(f"Written {OUT} ({OUT.stat().st_size / 1024:.0f} KB)", file=sys.stderr)


if __name__ == "__main__":
    build_document()
